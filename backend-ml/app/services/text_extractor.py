# app/services/text_extractor.py
"""
Giai đoạn [1]-[2] của pipeline parser: đọc file -> text thô -> chuẩn hóa.

Hỗ trợ 4 loại input:
  - PDF có text layer (đọc trực tiếp qua pdfplumber, nhanh, chính xác 100%)
  - PDF scan (ảnh, không có text layer) -> OCR từng trang qua Tesseract
  - DOCX -> đọc trực tiếp qua python-docx
  - Ảnh chụp/scan CV (JPG/PNG) -> OCR trực tiếp qua Tesseract

Cần cài Tesseract OCR (binary hệ thống, KHÔNG chỉ pip install pytesseract):
  Windows: https://github.com/UB-Mannheim/tesseract/wiki (chọn installer .exe)
  Sau khi cài, nếu không tự vào PATH, set đường dẫn thủ công — xem
  _TESSERACT_CMD_OVERRIDE bên dưới.
"""
from __future__ import annotations

import io
import os
import unicodedata

import pdfplumber
import pytesseract
from docx import Document
from PIL import Image

# Nếu Tesseract cài trên Windows nhưng không tự thêm vào PATH (thường gặp),
# set biến môi trường TESSERACT_CMD trỏ tới đường dẫn tesseract.exe, vd:
#   set TESSERACT_CMD=C:\Program Files\Tesseract-OCR\tesseract.exe
_TESSERACT_CMD_OVERRIDE = os.environ.get("TESSERACT_CMD")
if _TESSERACT_CMD_OVERRIDE:
    pytesseract.pytesseract.tesseract_cmd = _TESSERACT_CMD_OVERRIDE

# Ngôn ngữ OCR: chỉ tiếng Anh, vì CV đầu vào cố định là tiếng Anh.
_OCR_LANG = "eng"

_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png"}


class UnsupportedFileTypeError(Exception):
    pass


class OcrNotAvailableError(Exception):
    """Raise khi gọi OCR nhưng Tesseract binary chưa được cài trên máy."""
    pass


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, list[str]]:
    """
    Trả về (text, warnings).
    """
    warnings: list[str] = []
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        text = _extract_pdf(file_bytes, warnings)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    elif ext == "doc":
        raise UnsupportedFileTypeError(
            "File .doc (Word 97-2003) chưa được hỗ trợ, chỉ nhận .docx và .pdf"
        )
    elif ext in _IMAGE_EXTENSIONS:
        text = _extract_image(file_bytes, warnings)
    else:
        raise UnsupportedFileTypeError(f"Định dạng file không được hỗ trợ: .{ext}")

    return _normalize(text), warnings


def _ocr_image(image: Image.Image) -> str:
    try:
        return pytesseract.image_to_string(image, lang=_OCR_LANG)
    except pytesseract.TesseractNotFoundError as e:
        raise OcrNotAvailableError(
            "Tesseract OCR chưa được cài trên máy (chỉ pip install pytesseract "
            "là không đủ — cần cài binary hệ thống). Xem hướng dẫn ở đầu file "
            "text_extractor.py."
        ) from e


def _extract_pdf(file_bytes: bytes, warnings: list[str]) -> str:
    chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""

            if not page_text.strip():
                # Không có text layer -> khả năng cao là trang scan (ảnh) ->
                # rasterize trang thành ảnh rồi OCR. Dùng page.to_image() của
                # pdfplumber (dựa trên pypdfium2, đã bundle binary sẵn trong
                # wheel — KHÔNG cần cài thêm Poppler/ImageMagick riêng như
                # cách làm cũ với pdf2image).
                try:
                    page_image = page.to_image(resolution=200).original
                    ocr_text = _ocr_image(page_image)
                    if ocr_text.strip():
                        page_text = ocr_text
                        warnings.append(
                            f"Trang {i + 1}: không có text layer, đã OCR tự động "
                            f"(độ chính xác thấp hơn PDF text thật, nên kiểm tra lại)."
                        )
                    else:
                        warnings.append(
                            f"Trang {i + 1}: không có text layer VÀ OCR không đọc "
                            f"được nội dung — trang có thể trống hoặc ảnh quá mờ."
                        )
                except OcrNotAvailableError as e:
                    warnings.append(f"Trang {i + 1}: {e}")

            chunks.append(page_text)
    return "\n".join(chunks)


def _extract_image(file_bytes: bytes, warnings: list[str]) -> str:
    try:
        image = Image.open(io.BytesIO(file_bytes))
    except Exception as e:
        raise UnsupportedFileTypeError(f"Không đọc được file ảnh: {e}") from e

    try:
        text = _ocr_image(image)
    except OcrNotAvailableError as e:
        warnings.append(str(e))
        return ""

    if not text.strip():
        warnings.append(
            "OCR không đọc được nội dung nào từ ảnh — ảnh có thể quá mờ, "
            "nghiêng, hoặc độ phân giải quá thấp."
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs]
    # Nhiều CV dùng table để layout 2 cột -> phải quét cả table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _normalize(text: str) -> str:
    # Chuẩn hóa unicode tiếng Việt về dạng NFC (tránh lỗi dấu bị tách rời)
    text = unicodedata.normalize("NFC", text)
    # Gom nhiều dòng trống liên tiếp
    lines = [line.rstrip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line.strip() != "" or True)
    return cleaned.strip()